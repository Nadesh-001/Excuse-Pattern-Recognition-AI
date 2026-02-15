import csv
import io
import os
from flask import make_response, current_app
from repository.db import execute_query
from docx import Document
from fpdf import FPDF

def generate_csv_report(user_id=None, role='employee', type='tasks'):
    """Generates a CSV report with Excel-compatible UTF-8 BOM."""
    headers, data = _fetch_report_data(user_id, role, type)
    
    si = io.StringIO()
    si.write('\ufeff') # UTF-8 BOM
    cw = csv.writer(si)
    cw.writerow(headers)
    for row in data:
        cw.writerow(list(row.values()) if isinstance(row, dict) else list(row))
        
    response = make_response(si.getvalue())
    response.headers["Content-Disposition"] = f"attachment; filename={type}_report.csv"
    response.headers["Content-type"] = "text/csv; charset=utf-8-sig"
    return response

def generate_word_report(user_id=None, role='employee', type='tasks'):
    """Generates a Word (.docx) report."""
    headers, data = _fetch_report_data(user_id, role, type)
    
    doc = Document()
    doc.add_heading(f'Excuse Pattern AI - {type.capitalize()} Report', 0)
    
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = 'Table Grid'
    hdr_cells = table.rows[0].cells
    for i, header in enumerate(headers):
        hdr_cells[i].text = header
        
    for row in data:
        row_cells = table.add_row().cells
        row_values = list(row.values()) if isinstance(row, dict) else list(row)
        for i, val in enumerate(row_values):
            row_cells[i].text = str(val)
            
    target_stream = io.BytesIO()
    doc.save(target_stream)
    target_stream.seek(0)
    
    response = make_response(target_stream.getvalue())
    response.headers["Content-Disposition"] = f"attachment; filename={type}_report.docx"
    response.headers["Content-type"] = "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
    return response

def generate_pdf_report(user_id=None, role='employee', type='tasks'):
    """Generates a PDF report."""
    headers, data = _fetch_report_data(user_id, role, type)
    
    pdf = FPDF(orientation='L', unit='mm', format='A4')
    pdf.add_page()
    pdf.set_font("Helvetica", 'B', 16)
    pdf.cell(0, 10, f'Excuse Pattern AI - {type.capitalize()} Report', ln=True, align='C')
    pdf.ln(10)
    
    # Table headers
    pdf.set_font("Helvetica", 'B', 10)
    col_width = pdf.epw / len(headers)
    for header in headers:
        pdf.cell(col_width, 10, header, border=1)
    pdf.ln()
    
    # Table data
    pdf.set_font("Helvetica", size=8)
    for row in data:
        row_values = list(row.values()) if isinstance(row, dict) else list(row)
        for val in row_values:
            pdf.cell(col_width, 8, str(val)[:30], border=1) # Truncate for A4 width
        pdf.ln()
        
    response = make_response(pdf.output())
    response.headers["Content-Disposition"] = f"attachment; filename={type}_report.pdf"
    response.headers["Content-type"] = "application/pdf"
    return response

def _fetch_report_data(user_id, role, type):
    """Internal helper to fetch data for all formats."""
    data = []
    headers = []
    try:
        if type == 'tasks':
            where_clause = ""
            params = []
            if role == 'employee' and user_id:
                where_clause = "WHERE assigned_to = %s"
                params = [user_id]
            data = execute_query(f"SELECT id, title, status, priority, deadline FROM tasks {where_clause} ORDER BY created_at DESC", params)
            headers = ['ID', 'Title', 'Status', 'Priority', 'Deadline']
        else:
            where_clause = ""
            params = []
            if role == 'employee' and user_id:
                where_clause = "WHERE d.user_id = %s"
                params = [user_id]
            data = execute_query(f"""
                SELECT d.id, t.title, d.score_authenticity, d.risk_level, d.submitted_at 
                FROM delays d LEFT JOIN tasks t ON d.task_id = t.id 
                {where_clause} ORDER BY d.submitted_at DESC
            """, params)
            headers = ['ID', 'Task', 'Auth Score', 'Risk', 'Date']
    except Exception as e:
        current_app.logger.error(f"Export data fetch error: {e}")
    
    return headers, (data if data else [])
